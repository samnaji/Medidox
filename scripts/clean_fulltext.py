import sys
import os
from docx import Document
from docx.shared import Pt
import docx
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
import shutil
import glob
from textblob import TextBlob
import re
import docx
import os
from docx import Document
from docx.shared import Pt
import docx
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from transformers import AutoTokenizer, T5ForConditionalGeneration
import pdb


dir_in = sys.argv[1]
dir_out = sys.argv[2]
Archive=sys.argv[3]

Read_from_parent_dir=dir_in
Save_to_parent_dir=dir_out





def remove_non_english_characters(raw_input):
  # Keep the characters in the identifier and remove all other non-english characters
  clean_input = re.sub(r'[^\x00-\x7F]+','', raw_input).strip()
  clean_input = re.sub(r'\s+', ' ', clean_input)

  return clean_input


def correct_sentence_spelling(sentence):
    sentence = TextBlob(sentence)
    result = sentence.correct()
    return result.raw

import re
def clean_para_text(text): 
  out=correct_sentence_spelling(text)
  return out

def count_words(text):
    words = text.split(' ')
    count = 0
    for word in words:
        count += 1
    return count  


"""## Reading Methods"""
def check_header(text):
  date_formats = [ 
    r"\d\d/\d\d/\d\d\d\d",  # DD/MM/YYYY (e.g. "12/12/2022")
    r"\d\d/\d\d/\d\d",  # DD/MM/YY_ (e.g. "12/12/2022")
    r"\d\d/\d\d/\d\d",       # DD/MM/YY (e.g. "12/12/22")
    r"\d\d-\d\d-\d\d\d\d",  # DD-MM-YYYY (e.g. "12-12-2022")
    r"\d\d-\d\d-\d\d",  # DD-MM-YYYY (e.g. "12-12-22")
    r"\d\d\d\d-\d\d-\d\d",  # YYYY-MM-DD (e.g. "2022-12-12")
    r"\d\d\d\d/\d\d/\d\d",  # YYYY/MM/DD (e.g. "2022/12/12")
    r"\d\d-\w\w\w-\d\d\d\d", # DD-Mon-YYYY (e.g. "12-Dec-2022")
    r"\d\d \w\w\w+ \d\d", # DD Month YYYY (e.g. "12 December 2022")
    r"\w\w\w \d\d \d\d\d\d", # Mon DD YYYY (e.g. "Dec 12 2022")
    r"\d\d \w\w\w+ \d\d\d\d", # DD Month YYYY (e.g. "12 December 2022")
    r"\d\d \w\w\w+ \d\d", # DD Month YYYY (e.g. "12 December 22")
    r"\d\d/\d\d/\d\d\d\d",  # DD/MM/YYYY (e.g. "12/12/2022")
    r"\d\d/\d\d/\d\d",  # DD/MM/YY (e.g. "12/12/22")
    r"\d\d-\d\d-\d\d\d\d",  # DD-MM-YYYY (e.g. "12-12-2022")
    r"\d\d-\d\d-\d\d",  # DD-MM-YY (e.g. "12-12-22")
    r"\d\d\d\d-\d\d-\d\d",  # YYYY-MM-DD (e.g. "2022-12-12")
    r"\d\d\d\d/\d\d/\d\d",  # YYYY/MM/DD (e.g. "2022/12/12")
    r"\d\d-\w\w\w-\d\d\d\d", # DD-Mon-YYYY (e.g. "12-Dec-2022")
    r"\d\d \w\w\w+ \d\d\d\d", # DD Month YYYY (e.g. "12 December 2022")
    r"\w\w\w \d\d, \d\d\d\d", # Mon DD, YYYY (e.g. "Dec 12, 2022")
    r"\d\d \w\w\w \d\d\d\d", # DD Month YYYY (e.g. "12 December 2022")
    r"\d\d \w\w\w \d\d", # DD Month YY (e.g. "12 December 22")
    r"\d\d/\d\d/\d\d\d\d",  # MM/DD/YYYY (e.g. "12/12/2022")
    r"\d\d/\d\d/\d\d",  # MM/DD/YY (e.g. "12/12/22")
    r"\d\d-\d\d-\d\d\d\d",  # MM-DD-YYYY (e.g. "12-12-2022")
    r"\d\d-\d\d-\d\d",  # MM-DD-YY (e.g. "12-12-22")
    r"\d\d\d\d/\d\d/\d\d",  # YYYY/MM/DD (e.g. "2022/12/12")
    r"\d\d\d\d-\d\d-\d\d",  # YYYY-MM-DD (e.g. "2022-12-12")
    r"\d\d-\w\w\w-\d\d\d\d", # MM-Mon-YYYY (e.g. "12-Dec-2022")
    r"\d\d \w\w\w \d\d\d\d", # MM Month YYYY (e.g. "12 December 2022")
    r"\w\w\w \d\d, \d\d\d\d", # Mon DD, YYYY (e.g. "Dec 12, 2022")
    r"\d\d.\d\d.\d\d\d\d",  # MM.DD.YYYY (e.g. "12.12.2022")
    r"\d\d.\d\d.\d\d",  # MM.DD.YY (e.g. "12.12.22")
    r"\d\d \d\d \d\d\d\d",  # MM DD YYYY (e.g. "12 12 2022")
    r"\d\d.\d\d.\d\d\d\d",  # DD.MM.YYYY (e.g. "12.12.2022")
    r"\d\d/\d\d/\d\d\d\d",  # DD/MM/YYYY (e.g. "12/12/2022")
    r"\d\d/\d\d/\d\d",  # DD/MM/YY (e.g. "12/12/22")
    r"\d\d-\d\d-\d\d\d\d",  # DD-MM-YYYY (e.g. "12-12-2022")
    r"\d\d-\d\d-\d\d",  # DD-MM-YY (e.g. "12-12-22")
    r"\d\d\d\d/\d\d/\d\d",  # YYYY/DD/
    ]
  for pattern in date_formats:
    pattern=pattern+"_"
    match = re.search(pattern, text)
    if match:
      return match.group()
    return -1 #not found

def find_names(text):
  text = re.sub(r'[\d,._-]', '', text)
  text.split()
  out=text.split()
  return out[0:2]

"""## Auto Run"""

files = glob.glob(os.path.join(Read_from_parent_dir, '*.doc*'))

"""## Inference Model"""

files = glob.glob(os.path.join(Read_from_parent_dir, '*.doc*'))
total_number_docs=len(files)
max_length=512
for document_number in  range(total_number_docs):
 try:
  file_path=files[document_number]

  file_name=os.path.basename(file_path)
  Save_to_dir=os.path.join(Save_to_parent_dir,file_name)

  document = docx.Document(file_path)

  out_gram=[]
  out_paraph=[]
  out_type=[]
  original_text=[]
  # Loop through each paragraph in the document
  i=0
  Total=len(document.paragraphs)
  
  for paragraph in document.paragraphs:
    raw_input=paragraph.text
    # ** look this market 
    clean_input = remove_non_english_characters(raw_input) #
    if len(clean_input)==0:
      i+=1
      print(f'Document Number {document_number+1} out of {total_number_docs}: {i} out of {Total} blocks completed: Skipped')
      out_paraph.append('')
      original_text.append(clean_input)
      out_type.append('Skipped')
      continue
    if len(clean_input)>0:
      if len(clean_input.strip())>=15:
        date_possible_loc=clean_input[:15].strip()
      else:
        date_possible_loc=clean_input

      if not(check_header(date_possible_loc)== -1):
         
        date=check_header(clean_input) # header
        #out_text.append( [date]+process_header(clean_input))
        #out_gram.append(clean_input)
        out_paraph.append(raw_input)
        out_type.append("header")
        original_text.append(raw_input)
        i+=1
      else: 
        TEXT=clean_para_text(clean_input)
        out_paraph.append(TEXT) 
        #print(TEXT)
        #pdb.set_trace()
        out_type.append("text")
        original_text.append(raw_input)
        i+=1
    if len(out_type)<2:
      print(f'Document Number {document_number+1} out of {total_number_docs}: {i} out of {Total} blocks completed')
    else:
      print(f'Document Number {document_number+1} out of {total_number_docs}: {i} out of {Total} blocks completed: Block: {out_type[-1]}')



  # Create a new document
  document = Document()
  
  # Add a paragraph to the header
  paragraph = document.add_paragraph()
  para = document.add_paragraph(style='Normal')
  run = para.add_run("REVIEW OF RECORDS")
  font = run.font
  font.size = Pt(11)
  font.name = 'Courier New'
  font.bold = True
  para.alignment = WD_ALIGN_PARAGRAPH.CENTER
  names=find_names(file_name)
  for p in range(len(out_type)):
    if out_type[p]=="text":

        # Add paraphrase only
        paragraph = document.add_paragraph()
        para = document.add_paragraph(style='Normal')

        original_words = original_text[p].split()

        filtered_original_words = [re.sub(r'[^\w\s]|\s+', '', word).lower() for word in original_words]
        
        paraphrase_words = out_paraph[p].split()
        #filtered_paraphrase_words = [word.strip("-_^!?:;,. '\"").lower() for word in paraphrase_words]
        filtered_paraphrase_words = [re.sub(r'[^\w\s]|\s+', '', word).lower() for word in paraphrase_words]
        for i in range(len(filtered_paraphrase_words)):
            word=filtered_paraphrase_words[i]
            if i>0:
              word_out=" "+paraphrase_words[i]
            else:
              word_out=paraphrase_words[i]

            if word in names:
                run = para.add_run(word_out)
                font = run.font
                font.size = Pt(11)
                font.name = 'Courier New'
                font.highlight_color = WD_COLOR_INDEX.YELLOW  # set highlight color to green
            elif not word in filtered_original_words:
                run = para.add_run(word_out)
                font = run.font
                font.size = Pt(11)
                font.name = 'Courier New'
                font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN  # set highlight color to green
                #font.color.rgb = RGBColor(0, 0, 0)  # set font color to default
            else:
                run = para.add_run(word_out)
                font = run.font
                font.size = Pt(11)
                font.name = 'Courier New'
                #font.color.rgb = RGBColor(0, 0, 0)  # set font color to default

        font.size = Pt(11)
        font.name = 'Courier New'
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        #font.italic = False 
 

    
    elif out_type[p]=="header":
      # Add a paragraph to the document
      paragraph = document.add_paragraph()
      para = document.add_paragraph(style='Normal')
      run = para.add_run(out_paraph[p])
      font = run.font
      font.size = Pt(11)
      #font.name = 'Courier New'
      para.alignment = WD_ALIGN_PARAGRAPH.LEFT
      font.bold = True


  # Save the document
  document.save(Save_to_dir)
  print(f"saved: {file_name}")
  shutil.move(file_path, Archive)
 except Exception as e:
        # Exception handling code goes here
        print(f"An error occurred: {e}") 


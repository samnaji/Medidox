import sys
import os
from docx import Document
from docx.shared import Pt
import docx
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
import shutil
import glob
from textblob import TextBlob
import re
import docx
import os
from docx import Document
from docx.shared import Pt
import docx
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from transformers import AutoTokenizer, T5ForConditionalGeneration
import pdb


dir_in = sys.argv[1]
dir_out = sys.argv[2]
dir_train = sys.argv[3]
dir_model = sys.argv[4]
num_outputs = int(sys.argv[5])
num_v = int(sys.argv[6])
Marker_input = sys.argv[7]

Read_from_parent_dir=dir_in
Save_to_parent_dir=dir_out
Move_to_train=dir_train
Load_from_models_dir=dir_model
Model_Selected=num_v# from front end
Marker_input=Marker_input  # *** Connect this to input by user





def paraphrase_med(text,max_length):
    text="paraphrase: "+text
    inputs = tokenizer(text, padding='longest', max_length=max_length, return_tensors='pt')
    input_ids = inputs.input_ids
    attention_mask = inputs.attention_mask
    output = model.generate(input_ids, attention_mask=attention_mask, max_length=max_length)
    return tokenizer.decode(output[0], skip_special_tokens=True)

def remove_non_english_characters(raw_input, identifier):
  # Keep the characters in the identifier and remove all other non-english characters
  if identifier == -1:
    clean_input = re.sub(rf'[^{identifier}\x00-\x7F]+','', raw_input).strip()
    clean_input = re.sub(r'\s+', ' ', clean_input)

  else:
    clean_input = re.sub(r'[^\x00-\x7F]+','', raw_input).strip()
    clean_input = re.sub(r'\s+', ' ', clean_input)

  return clean_input


def correct_sentence_spelling(sentence):
    sentence = TextBlob(sentence)
    result = sentence.correct()
    return result.raw

import re
def clean_para_text(text,marker,paraphrase_med,max_length): 
  indices = [m.start() for m in re.finditer(marker, text)]
  if len(indices)==0 or Marker_input=="":
    return run_inference_on_text(text, paraphrase_med,max_length)
  if not len(indices)%2==0:
    raise Exception(f"Unbalanced markers not allowed")
  pairs = list(zip(indices[::2], indices[1::2]))
  out=[]
  out_type=[]
  prev=0
  for i, (start, stop) in enumerate(pairs):
    substring=text[prev:start]
    if count_words(substring)>10:
      #segments.append((substring, "Paraphrase"))
      out.append(run_inference_on_text(substring, paraphrase_med,max_length))
    elif len(substring)>0:
      out.append(correct_sentence_spelling(substring))

    #clean
    substring=text[start+1:stop]
    if len(substring)>0:
      if start>0:
        out.append(" "+substring)
      else:
        out.append(" "+substring)

    prev=stop+1
  #last_substring 

  if prev<len(text):
    substring=text[stop+1:]   
    if len(substring)<10:
      out.append(correct_sentence_spelling(substring))
    else:
      out.append(run_inference_on_text(substring, paraphrase_med,max_length))
  return ' '.join(out)

def count_words(text):
    words = text.split(' ')
    count = 0
    for word in words:
        count += 1
    return count  

def run_inference_on_text(text, InferencFunc,max_length):
  P_a=0
  out=[]
  results=[]
  while not(text==""):
    END=len(text)
    if END<=max_length:
        out.append(text)
        P_a=len(text)
        text=text[P_a:]
        break
    else:
        P_dot=text[:max_length].rfind(".")
        P_ques=text[:max_length].rfind("?")
        P_ex=text[:max_length].rfind("!")
        P_a=max([P_dot,P_ques,P_ex])
        if P_a==-1: #if  .!? is not found look for 
          P_a=text[:max_length].rfind(",")
          if P_a==-1: #if  , is not found look for 
            P_a=max_length
            out.append(text[:P_a+1])
            text=text[P_a+1:]

          else: #if  , is  found look for 
            out.append(text[:P_a+1])
            text=text[P_a+1:]

        else: #if  . ? or ! are found look for 
          out.append(text[:P_a+1])
          text=text[P_a+1:]

  for part in out:
    results.append(InferencFunc(part,max_length))
  return ''.join(results)

"""## Reading Methods"""
def check_header(text):
  date_formats = [ 
    r"\d\d/\d\d/\d\d\d\d",  # DD/MM/YYYY (e.g. "12/12/2022")
    r"\d\d/\d\d/\d\d",  # DD/MM/YY_ (e.g. "12/12/2022")
    r"\d\d/\d\d/\d\d",       # DD/MM/YY (e.g. "12/12/22")
    r"\d\d-\d\d-\d\d\d\d",  # DD-MM-YYYY (e.g. "12-12-2022")
    r"\d\d-\d\d-\d\d",  # DD-MM-YYYY (e.g. "12-12-22")
    r"\d\d\d\d-\d\d-\d\d",  # YYYY-MM-DD (e.g. "2022-12-12")
    r"\d\d\d\d/\d\d/\d\d",  # YYYY/MM/DD (e.g. "2022/12/12")
    r"\d\d-\w\w\w-\d\d\d\d", # DD-Mon-YYYY (e.g. "12-Dec-2022")
    r"\d\d \w\w\w+ \d\d", # DD Month YYYY (e.g. "12 December 2022")
    r"\w\w\w \d\d \d\d\d\d", # Mon DD YYYY (e.g. "Dec 12 2022")
    r"\d\d \w\w\w+ \d\d\d\d", # DD Month YYYY (e.g. "12 December 2022")
    r"\d\d \w\w\w+ \d\d", # DD Month YYYY (e.g. "12 December 22")
    r"\d\d/\d\d/\d\d\d\d",  # DD/MM/YYYY (e.g. "12/12/2022")
    r"\d\d/\d\d/\d\d",  # DD/MM/YY (e.g. "12/12/22")
    r"\d\d-\d\d-\d\d\d\d",  # DD-MM-YYYY (e.g. "12-12-2022")
    r"\d\d-\d\d-\d\d",  # DD-MM-YY (e.g. "12-12-22")
    r"\d\d\d\d-\d\d-\d\d",  # YYYY-MM-DD (e.g. "2022-12-12")
    r"\d\d\d\d/\d\d/\d\d",  # YYYY/MM/DD (e.g. "2022/12/12")
    r"\d\d-\w\w\w-\d\d\d\d", # DD-Mon-YYYY (e.g. "12-Dec-2022")
    r"\d\d \w\w\w+ \d\d\d\d", # DD Month YYYY (e.g. "12 December 2022")
    r"\w\w\w \d\d, \d\d\d\d", # Mon DD, YYYY (e.g. "Dec 12, 2022")
    r"\d\d \w\w\w \d\d\d\d", # DD Month YYYY (e.g. "12 December 2022")
    r"\d\d \w\w\w \d\d", # DD Month YY (e.g. "12 December 22")
    r"\d\d/\d\d/\d\d\d\d",  # MM/DD/YYYY (e.g. "12/12/2022")
    r"\d\d/\d\d/\d\d",  # MM/DD/YY (e.g. "12/12/22")
    r"\d\d-\d\d-\d\d\d\d",  # MM-DD-YYYY (e.g. "12-12-2022")
    r"\d\d-\d\d-\d\d",  # MM-DD-YY (e.g. "12-12-22")
    r"\d\d\d\d/\d\d/\d\d",  # YYYY/MM/DD (e.g. "2022/12/12")
    r"\d\d\d\d-\d\d-\d\d",  # YYYY-MM-DD (e.g. "2022-12-12")
    r"\d\d-\w\w\w-\d\d\d\d", # MM-Mon-YYYY (e.g. "12-Dec-2022")
    r"\d\d \w\w\w \d\d\d\d", # MM Month YYYY (e.g. "12 December 2022")
    r"\w\w\w \d\d, \d\d\d\d", # Mon DD, YYYY (e.g. "Dec 12, 2022")
    r"\d\d.\d\d.\d\d\d\d",  # MM.DD.YYYY (e.g. "12.12.2022")
    r"\d\d.\d\d.\d\d",  # MM.DD.YY (e.g. "12.12.22")
    r"\d\d \d\d \d\d\d\d",  # MM DD YYYY (e.g. "12 12 2022")
    r"\d\d.\d\d.\d\d\d\d",  # DD.MM.YYYY (e.g. "12.12.2022")
    r"\d\d/\d\d/\d\d\d\d",  # DD/MM/YYYY (e.g. "12/12/2022")
    r"\d\d/\d\d/\d\d",  # DD/MM/YY (e.g. "12/12/22")
    r"\d\d-\d\d-\d\d\d\d",  # DD-MM-YYYY (e.g. "12-12-2022")
    r"\d\d-\d\d-\d\d",  # DD-MM-YY (e.g. "12-12-22")
    r"\d\d\d\d/\d\d/\d\d",  # YYYY/DD/
    ]
  for pattern in date_formats:
    pattern=pattern+"_"
    match = re.search(pattern, text)
    if match:
      return match.group()
    return -1 #not found
      
def process_header(text):
  parts = re.split("[,]", text)
  names = []
  p=0
  for part in parts:
    if part:
      p+=1
      # add the part to the list of names
      if p==1:
        part=part[1:].lstrip()
      if p==2: # doctor
        part = part.lstrip()
        if part[:2]=="MD":
          part=part[2:].lstrip() + ", MD"
      names.append(part.lstrip())
  
  if len(parts)<2:
    headers[1]="Unknown" #MD

  headers[0]= names[0]
  headers[1]=names[1] #MD
  headers[2]= ", ".join(names[2:]) #other info
  return headers

def find_names(text):
  text = re.sub(r'[\d,._-]', '', text)
  text.split()
  out=text.split()
  return out[0:2]

"""## Auto Run"""
# Get a list of all subdirectories in the directory
models_folders =[entry.name for entry in os.scandir(Load_from_models_dir) if entry.is_dir() and entry.name.startswith('T5-')]
# Sort the list of subdirectories by the date they were created
models_folders.sort(key=lambda x: os.path.getctime(os.path.join(Load_from_models_dir, x)), reverse=True)
# Print the list of subdirectories
model_dir=os.path.join(Load_from_models_dir,models_folders[Model_Selected])
print(f"loaded: {model_dir}")
files = glob.glob(os.path.join(Read_from_parent_dir, '*.doc*'))

"""## Inference Model"""
CKPT=model_dir
print("model loaded:" +model_dir)
tokenizer = AutoTokenizer.from_pretrained(CKPT)
model = T5ForConditionalGeneration.from_pretrained(CKPT)

files = glob.glob(os.path.join(Read_from_parent_dir, '*.doc*'))
total_number_docs=len(files)
max_length=512
for document_number in  range(total_number_docs):
 try:
  file_path=files[document_number]

  file_name=os.path.basename(file_path)
  Save_to_dir=os.path.join(Save_to_parent_dir,file_name)

  document = docx.Document(file_path)

  out_gram=[]
  out_paraph=[]
  out_type=[]
  original_text=[]
  # Loop through each paragraph in the document
  i=0
  Total=len(document.paragraphs)
  
  for paragraph in document.paragraphs:
    raw_input=paragraph.text
    # ** look this market 
    clean_input = remove_non_english_characters(raw_input, Marker_input) #
    if len(clean_input)==0:
      i+=1
      print(f'Document Number {document_number}: {i} out of {Total} completed: Skipped')
      out_paraph.append('')
      original_text.append(clean_input)
      out_type.append('Skipped')
      continue
    if len(clean_input)>0:
      if len(clean_input.strip())>=15:
        date_possible_loc=clean_input[:15].strip()
      else:
        date_possible_loc=clean_input

      if not(check_header(date_possible_loc)== -1):
         
        date=check_header(clean_input) # header
        #out_text.append( [date]+process_header(clean_input))
        #out_gram.append(clean_input)
        out_paraph.append(raw_input)
        out_type.append("header")
        original_text.append(raw_input)
        i+=1
      else: 
        TEXT=clean_para_text(clean_input,Marker_input,paraphrase_med,max_length)
        out_paraph.append(TEXT) 
        #print(TEXT)
        #pdb.set_trace()
        out_type.append("text")
        original_text.append(raw_input)
        i+=1
    if len(out_type)<2:
      print(f'Document Number {document_number}: {i} out of {Total} completed')
    else:
      print(f'Document Number {document_number}: {i} out of {Total} completed: Block: {out_type[-1]}')



  # Create a new document
  document = Document()
  
  # Add a paragraph to the header
  paragraph = document.add_paragraph()
  para = document.add_paragraph(style='Normal')
  run = para.add_run("REVIEW OF RECORDS")
  font = run.font
  font.size = Pt(11)
  font.name = 'Courier New'
  font.bold = True
  para.alignment = WD_ALIGN_PARAGRAPH.CENTER
  names=find_names(file_name)
  for p in range(len(out_type)):
    if out_type[p]=="text":

        # Add paraphrase only
        paragraph = document.add_paragraph()
        para = document.add_paragraph(style='Normal')

        original_words = re.sub(rf'[{Marker_input}]', ' ',original_text[p]).split()

        filtered_original_words = [re.sub(r'[^\w\s]|\s+', '', word).lower() for word in original_words]
        
        paraphrase_words = out_paraph[p].split()
        #filtered_paraphrase_words = [word.strip("-_^!?:;,. '\"").lower() for word in paraphrase_words]
        filtered_paraphrase_words = [re.sub(r'[^\w\s]|\s+', '', word).lower() for word in paraphrase_words]
        for i in range(len(filtered_paraphrase_words)):
            word=filtered_paraphrase_words[i]
            if i>0:
              word_out=" "+paraphrase_words[i]
            else:
              word_out=paraphrase_words[i]

            if word in names:
                run = para.add_run(word_out)
                font = run.font
                font.size = Pt(11)
                font.name = 'Courier New'
                font.highlight_color = WD_COLOR_INDEX.YELLOW  # set highlight color to green
            elif not word in filtered_original_words:
                run = para.add_run(word_out)
                font = run.font
                font.size = Pt(11)
                font.name = 'Courier New'
                font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN  # set highlight color to green
                #font.color.rgb = RGBColor(0, 0, 0)  # set font color to default
            else:
                run = para.add_run(word_out)
                font = run.font
                font.size = Pt(11)
                font.name = 'Courier New'
                #font.color.rgb = RGBColor(0, 0, 0)  # set font color to default

        font.size = Pt(11)
        font.name = 'Courier New'
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        #font.italic = False 
 

    
    elif out_type[p]=="header":
      # Add a paragraph to the document
      paragraph = document.add_paragraph()
      para = document.add_paragraph(style='Normal')
      run = para.add_run(out_paraph[p])
      font = run.font
      font.size = Pt(11)
      #font.name = 'Courier New'
      para.alignment = WD_ALIGN_PARAGRAPH.LEFT
      font.bold = True


  # Save the document
  document.save(Save_to_dir)
  print(f"saved: {file_name}")
  shutil.move(file_path, Move_to_train)
 except Exception as e:
        # Exception handling code goes here
        print(f"An error occurred: {e}") 

